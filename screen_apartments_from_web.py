#!/usr/bin/env python3
"""Screen apartment websites for unit, price, laundry, parking, and pet signals.

This is intentionally conservative: it extracts public text from official pages,
records short evidence snippets, and marks uncertain fields instead of guessing.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import html
import json
import re
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo


DEFAULT_INPUT = "apartments_google_budget_clean.csv"
DEFAULT_OUTPUT_CSV = "apartment_screening_auto.csv"
DEFAULT_OUTPUT_XLSX = "apartment_screening_auto.xlsx"
DEFAULT_OUTPUT_DOCX = "apartment_screening_auto_summary.docx"
DEFAULT_CACHE_DIR = ".web_cache"
DEFAULT_OVERRIDES = "manual_overrides.json"

UNIT_RE = re.compile(
    r"\b(studio|bachelor|junior\s*-?\s*1|junior\s+one|[1-6]\s*(?:b|x|bed|beds|bedroom|bedrooms|br)\s*[/-]?\s*[1-6]?\s*(?:b|bath|baths|bathroom|bathrooms|ba)?|[1-6]\s*bed(?:room)?(?:\s*/\s*[1-6]\s*bath(?:room)?)?|one\s*bed(?:room)?)\b",
    re.IGNORECASE,
)
PRICE_RE = re.compile(r"\$\s?([1-9]\d{2,4})(?:[,\d]{0,4})?(?:\s?(?:-|to)\s?\$?\s?([1-9]\d{2,4}))?", re.IGNORECASE)
SQFT_RE = re.compile(r"\b(\d{3,4})\s*(?:-|to)?\s*(\d{3,4})?\s*(?:sq\.?\s*ft\.?|square\s*feet|sf)\b", re.IGNORECASE)
LAUNDRY_RE = re.compile(r"\b(washer|dryer|laundry|washers|dryers|w/d|in-unit|in unit|in-home|in home)\b", re.IGNORECASE)
PARKING_RE = re.compile(r"\b(parking|garage|assigned|covered|gated|tandem|ev charging|reserved)\b", re.IGNORECASE)
PET_RE = re.compile(r"\b(pet|pets|cat|dog|animal)\b", re.IGNORECASE)
BATH_RE = re.compile(r"\b(private bath(?:room)?|private ensuite|en[- ]?suite|own bath(?:room)?|shared bath(?:room)?|bath(?:room)?|[1-6]\s*bath(?:room)?|[1-6]ba)\b", re.IGNORECASE)
PHONE_RE = re.compile(r"(?:\+?1[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}")
CALL_RE = re.compile(r"\b(call|contact|phone|leasing office|schedule a tour|details|pricing)\b", re.IGNORECASE)
PER_PERSON_RE = re.compile(r"\b(per person|/person|per bed|/bed|per bedroom|individual lease|by the bed|bedspace|roommate matching)\b", re.IGNORECASE)
WHOLE_UNIT_RE = re.compile(r"\b(per unit|whole unit|entire unit|entire apartment|full apartment|monthly rent|rent/month|apartment rent)\b", re.IGNORECASE)

BAD_PRICE_CONTEXT = re.compile(r"\b(application|admin|deposit|fee|pet|parking|garage|utility|utilities|income|holding)\b", re.IGNORECASE)
NEGATIVE_IN_UNIT_RE = re.compile(r"\b(no|not|without|shared|community|common|laundry room|facility|facilities)\b", re.IGNORECASE)
POSITIVE_IN_UNIT_RE = re.compile(r"\b(in-unit|in unit|in-home|in home|washer/dryer in|w/d in|washer and dryer in)\b", re.IGNORECASE)

LINK_KEYWORDS = (
    "floor",
    "plan",
    "availability",
    "available",
    "apartment",
    "amenit",
    "faq",
    "parking",
    "pet",
    "contact",
    "lease",
    "gallery",
)


class TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self.links: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style", "noscript", "svg"}:
            self._skip_depth += 1
            return
        if tag == "a":
            href = dict(attrs).get("href")
            if href:
                self.links.append(href)
        if tag in {"br", "p", "div", "li", "tr", "h1", "h2", "h3", "h4"}:
            self.parts.append(" ")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript", "svg"} and self._skip_depth:
            self._skip_depth -= 1
        if tag in {"p", "div", "li", "tr", "h1", "h2", "h3", "h4"}:
            self.parts.append(" ")

    def handle_data(self, data: str) -> None:
        if not self._skip_depth:
            self.parts.append(data)

    def text(self) -> str:
        raw = html.unescape(" ".join(self.parts))
        return re.sub(r"\s+", " ", raw).strip()


@dataclass
class PageData:
    url: str
    text: str
    links: list[str] = field(default_factory=list)
    error: str = ""


@dataclass
class ScreenResult:
    name: str
    group: str
    address: str
    distance_meters: str
    website: str
    target_unit: str = "Unknown"
    matched_target: str = "No"
    price_text: str = "Unknown"
    price_min: int | None = None
    price_basis: str = "Unknown"
    sqft_text: str = "Unknown"
    private_bath: str = "Unknown"
    laundry: str = "Unknown"
    parking: str = "Unknown"
    pet_policy: str = "Not ranked"
    phone: str = "Unknown"
    fit_status: str = "Needs review"
    confidence: str = "Low"
    evidence: str = ""
    pages_checked: str = ""
    notes: str = ""


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Extract apartment screening fields from website text.")
    parser.add_argument("--input", default=DEFAULT_INPUT, help="Clean apartment CSV from the Google search step.")
    parser.add_argument("--output-csv", default=DEFAULT_OUTPUT_CSV)
    parser.add_argument("--output-xlsx", default=DEFAULT_OUTPUT_XLSX)
    parser.add_argument("--output-docx", default=DEFAULT_OUTPUT_DOCX)
    parser.add_argument("--cache-dir", default=DEFAULT_CACHE_DIR)
    parser.add_argument("--overrides", default=DEFAULT_OVERRIDES)
    parser.add_argument("--budget", type=int, default=2500)
    parser.add_argument(
        "--target-units",
        default="studio,1b1b",
        help="Comma-separated acceptable unit types, e.g. studio,1b1b,2b2b,3b3b.",
    )
    parser.add_argument(
        "--private-bath",
        choices=["any", "prefer", "require"],
        default="any",
        help="Whether private bathroom should affect sorting/status.",
    )
    parser.add_argument("--max-pages-per-site", type=int, default=5)
    parser.add_argument("--delay-seconds", type=float, default=0.4)
    parser.add_argument("--limit", type=int, default=0, help="For testing: only process the first N website rows.")
    parser.add_argument("--cache-only", action="store_true", help="Do not fetch network pages; use cache and overrides only.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    rows = read_website_rows(args.input)
    if args.limit:
        rows = rows[: args.limit]

    overrides = read_overrides(args.overrides)
    cache_dir = Path(args.cache_dir)
    cache_dir.mkdir(exist_ok=True)

    results = []
    for index, row in enumerate(rows, start=1):
        name = row.get("name", "").strip()
        print(f"[{index}/{len(rows)}] Screening {name}", file=sys.stderr)
        result = screen_row(
            row=row,
            cache_dir=cache_dir,
            overrides=overrides.get(normalize_name(name), {}),
            budget=args.budget,
            acceptable_units=parse_unit_preferences(args.target_units),
            private_bath_mode=args.private_bath,
            max_pages=args.max_pages_per_site,
            delay_seconds=args.delay_seconds,
            cache_only=args.cache_only,
        )
        results.append(result)

    acceptable_units = parse_unit_preferences(args.target_units)
    results.sort(key=lambda item: sort_key(item, args.budget, args.private_bath))
    write_csv(args.output_csv, results)
    write_xlsx(args.output_xlsx, results)
    write_docx(args.output_docx, results, args.budget, acceptable_units, args.private_bath)
    print(f"Wrote {args.output_csv}, {args.output_xlsx}, and {args.output_docx}")
    return 0


def read_website_rows(path: str) -> list[dict[str, str]]:
    with open(path, newline="", encoding="utf-8") as file:
        rows = list(csv.DictReader(file))
    return [row for row in rows if row.get("website_uri", "").strip()]


def read_overrides(path: str) -> dict[str, dict[str, object]]:
    override_path = Path(path)
    if not override_path.exists():
        return {}
    with open(override_path, encoding="utf-8") as file:
        raw = json.load(file)
    return {normalize_name(name): value for name, value in raw.items()}


def screen_row(
    row: dict[str, str],
    cache_dir: Path,
    overrides: dict[str, object],
    budget: int,
    acceptable_units: list[dict[str, int | str | None]],
    private_bath_mode: str,
    max_pages: int,
    delay_seconds: float,
    cache_only: bool,
) -> ScreenResult:
    result = ScreenResult(
        name=row.get("name", "").strip(),
        group=row.get("clean_group", "").strip(),
        address=row.get("address", "").strip(),
        distance_meters=row.get("straight_line_meters", "").strip(),
        website=row.get("website_uri", "").strip(),
    )

    pages = collect_pages(result.website, cache_dir, max_pages, delay_seconds, cache_only)
    combined_text = " ".join(page.text for page in pages if page.text)
    result.pages_checked = " | ".join(page.url for page in pages)

    if combined_text:
        extracted = extract_fields(combined_text, budget, acceptable_units)
        for key, value in extracted.items():
            setattr(result, key, value)
    else:
        errors = "; ".join(page.error for page in pages if page.error)
        result.notes = f"No page text extracted. {errors}".strip()

    apply_overrides(result, overrides)
    result.fit_status = fit_status(result, budget, private_bath_mode)
    result.confidence = confidence(result)
    return result


def collect_pages(url: str, cache_dir: Path, max_pages: int, delay_seconds: float, cache_only: bool) -> list[PageData]:
    first = fetch_page(url, cache_dir, cache_only)
    pages = [first]
    candidate_urls = discover_candidate_urls(url, first.links)

    for candidate in candidate_urls:
        if len(pages) >= max_pages:
            break
        if normalize_url(candidate) in {normalize_url(page.url) for page in pages}:
            continue
        if not cache_only:
            time.sleep(delay_seconds)
        pages.append(fetch_page(candidate, cache_dir, cache_only))
    return pages


def fetch_page(url: str, cache_dir: Path, cache_only: bool) -> PageData:
    cache_path = cache_dir / f"{hashlib.sha256(url.encode('utf-8')).hexdigest()}.html"
    html_text = ""
    error = ""

    if cache_path.exists():
        html_text = cache_path.read_text(encoding="utf-8", errors="ignore")
    elif cache_only:
        error = "cache missing"
    else:
        try:
            request = urllib.request.Request(
                url,
                headers={
                    "User-Agent": "Mozilla/5.0 apartment-screening-research/1.0",
                    "Accept": "text/html,application/xhtml+xml",
                },
            )
            with urllib.request.urlopen(request, timeout=25) as response:
                content_type = response.headers.get("content-type", "")
                body = response.read(2_000_000)
            if "text/html" not in content_type and "application/xhtml" not in content_type:
                error = f"non-html content: {content_type}"
            html_text = body.decode("utf-8", errors="ignore")
            cache_path.write_text(html_text, encoding="utf-8")
        except (urllib.error.URLError, TimeoutError, ValueError) as exc:
            error = str(exc)

    if not html_text:
        return PageData(url=url, text="", links=[], error=error)

    parser = TextExtractor()
    parser.feed(html_text)
    return PageData(url=url, text=parser.text(), links=parser.links, error=error)


def discover_candidate_urls(base_url: str, links: list[str]) -> list[str]:
    parsed_base = urllib.parse.urlparse(base_url)
    base_domain = parsed_base.netloc.replace("www.", "")
    result = []
    for link in links:
        absolute = urllib.parse.urljoin(base_url, link)
        parsed = urllib.parse.urlparse(absolute)
        if parsed.scheme not in {"http", "https"}:
            continue
        if parsed.netloc.replace("www.", "") != base_domain:
            continue
        searchable = f"{parsed.path} {parsed.query}".lower()
        if any(keyword in searchable for keyword in LINK_KEYWORDS):
            cleaned = urllib.parse.urlunparse((parsed.scheme, parsed.netloc, parsed.path, "", parsed.query, ""))
            result.append(cleaned)
    return unique_keep_order(result)


def parse_unit_preferences(value: str) -> list[dict[str, int | str | None]]:
    units = []
    for raw in value.split(","):
        parsed = parse_unit_label(raw)
        if parsed and parsed not in units:
            units.append(parsed)
    return units or [parse_unit_label("studio"), parse_unit_label("1b1b")]


def parse_unit_label(value: str) -> dict[str, int | str | None] | None:
    text = normalize_name(value)
    if not text:
        return None
    if "studio" in text or "bachelor" in text:
        return {"label": "Studio", "beds": 0, "baths": 1}
    if "junior" in text:
        return {"label": "Junior 1", "beds": 1, "baths": 1}
    if "one bed" in text:
        return {"label": "1B1B", "beds": 1, "baths": 1}
    numbers = [int(number) for number in re.findall(r"[1-6]", text)]
    if numbers:
        beds = numbers[0]
        baths = numbers[1] if len(numbers) > 1 else None
        label = f"{beds}B{baths}B" if baths is not None else f"{beds}B"
        return {"label": label, "beds": beds, "baths": baths}
    return None


def unit_matches_preferences(unit_label: str, acceptable_units: list[dict[str, int | str | None]]) -> bool:
    parsed = parse_unit_label(unit_label)
    if not parsed:
        return False
    for target in acceptable_units:
        if parsed["beds"] != target["beds"]:
            continue
        if target["baths"] is None or parsed["baths"] is None or parsed["baths"] == target["baths"]:
            return True
    return False


def extract_fields(text: str, budget: int, acceptable_units: list[dict[str, int | str | None]]) -> dict[str, object]:
    unit_contexts = contexts(text, UNIT_RE, radius=120)
    price_contexts = [ctx for ctx in contexts(text, PRICE_RE, radius=90) if not BAD_PRICE_CONTEXT.search(ctx)]
    sqft_contexts = contexts(text, SQFT_RE, radius=60)
    laundry_contexts = contexts(text, LAUNDRY_RE, radius=95)
    parking_contexts = contexts(text, PARKING_RE, radius=85)
    pet_contexts = contexts(text, PET_RE, radius=85)
    bath_contexts = contexts(text, BATH_RE, radius=95)
    call_contexts = contexts(text, CALL_RE, radius=95)

    unit_summary = summarize_units(unit_contexts)
    matched_target = summarize_matching_units(unit_summary, acceptable_units)
    price_values = extract_price_values(price_contexts, matched_target, acceptable_units)
    price_min = min([item["per_person"] for item in price_values], default=None)

    return {
        "target_unit": matched_target if matched_target != "Unknown" else unit_summary,
        "matched_target": "Yes" if matched_target != "Unknown" else "No",
        "price_text": summarize_price(price_values),
        "price_min": price_min,
        "price_basis": summarize_price_basis(price_values),
        "sqft_text": summarize_matches(sqft_contexts, SQFT_RE, "Unknown"),
        "private_bath": summarize_private_bath(bath_contexts, matched_target),
        "laundry": summarize_laundry(laundry_contexts),
        "parking": summarize_signal(parking_contexts, "Unknown"),
        "pet_policy": summarize_signal(pet_contexts, "Not ranked"),
        "phone": summarize_phone(text, call_contexts),
        "evidence": compact_evidence(unit_contexts + price_contexts[:3] + sqft_contexts[:2] + bath_contexts[:2] + laundry_contexts[:2] + parking_contexts[:2] + call_contexts[:2]),
        "notes": notes_for_extraction(price_min, budget, matched_target, laundry_contexts, call_contexts),
    }


def contexts(text: str, pattern: re.Pattern[str], radius: int) -> list[str]:
    result = []
    for match in pattern.finditer(text):
        start = max(0, match.start() - radius)
        end = min(len(text), match.end() + radius)
        result.append(clean_snippet(text[start:end]))
    return unique_keep_order(result)[:12]


def extract_price_values(
    price_contexts: list[str],
    matched_target: str,
    acceptable_units: list[dict[str, int | str | None]],
) -> list[dict[str, int | str]]:
    values = []
    target_beds = best_bed_count(matched_target, acceptable_units)
    for context in price_contexts:
        for match in PRICE_RE.finditer(context):
            for group in match.groups():
                if not group:
                    continue
                digits = re.sub(r"\D", "", group)
                if digits:
                    value = int(digits)
                    if 900 <= value <= 10000:
                        basis = price_basis_from_context(context)
                        per_person = normalize_price_to_per_person(value, basis, target_beds)
                        values.append(
                            {
                                "raw": value,
                                "per_person": per_person,
                                "basis": basis,
                                "context": clean_snippet(context)[:260],
                            }
                        )
    return values


def summarize_units(contexts_: list[str]) -> str:
    if not contexts_:
        return "Unknown"
    found = set()
    joined = " ".join(contexts_).lower()
    if "studio" in joined or "bachelor" in joined:
        found.add("Studio")
    if "junior" in joined:
        found.add("Junior 1")
    if re.search(r"\b(1\s*bed|one\s*bed|1b1b|1\s*br)\b", joined):
        found.add("1B1B")
    for match in re.finditer(r"\b([1-6])\s*(?:b|x|bed|beds|bedroom|bedrooms|br)\s*[/-]?\s*([1-6])?\s*(?:b|bath|baths|bathroom|bathrooms|ba)?\b", joined):
        beds = int(match.group(1))
        baths = int(match.group(2)) if match.group(2) else None
        found.add(f"{beds}B{baths}B" if baths else f"{beds}B")
    return ", ".join(sorted(found)) if found else "Target unit mentioned"


def summarize_matching_units(unit_summary: str, acceptable_units: list[dict[str, int | str | None]]) -> str:
    if unit_summary == "Unknown":
        return "Unknown"
    matches = []
    for unit in [item.strip() for item in unit_summary.split(",")]:
        if unit_matches_preferences(unit, acceptable_units):
            matches.append(unit)
    return ", ".join(unique_keep_order(matches)) if matches else "Unknown"


def summarize_price(values: list[dict[str, int | str]]) -> str:
    if not values:
        return "Unknown"
    per_person_values = sorted(set(int(item["per_person"]) for item in values))
    if len(per_person_values) == 1:
        return f"${per_person_values[0]:,} per person"
    return f"${per_person_values[0]:,}+ per person; found up to ${per_person_values[-1]:,}"


def summarize_price_basis(values: list[dict[str, int | str]]) -> str:
    if not values:
        return "Unknown"
    bases = unique_keep_order([str(item["basis"]) for item in values])
    return ", ".join(bases)


def price_basis_from_context(context: str) -> str:
    if PER_PERSON_RE.search(context):
        return "per person"
    if WHOLE_UNIT_RE.search(context):
        return "whole unit"
    return "unknown; treated as per person"


def normalize_price_to_per_person(value: int, basis: str, target_beds: int | None) -> int:
    if basis == "whole unit" and target_beds and target_beds > 1:
        return round(value / target_beds)
    return value


def best_bed_count(matched_target: str, acceptable_units: list[dict[str, int | str | None]]) -> int | None:
    if matched_target != "Unknown":
        parsed = parse_unit_label(matched_target.split(",")[0])
        if parsed:
            return int(parsed["beds"] or 1)
    for target in acceptable_units:
        beds = target.get("beds")
        if isinstance(beds, int):
            return max(1, beds)
    return None


def summarize_matches(contexts_: list[str], pattern: re.Pattern[str], default: str) -> str:
    matches = []
    for context in contexts_:
        for match in pattern.finditer(context):
            matches.append(match.group(0))
    return "; ".join(unique_keep_order(matches)[:4]) if matches else default


def summarize_laundry(contexts_: list[str]) -> str:
    if not contexts_:
        return "Unknown"
    joined = " ".join(contexts_).lower()
    if POSITIVE_IN_UNIT_RE.search(joined) and not NEGATIVE_IN_UNIT_RE.search(joined):
        return "Likely in-unit washer/dryer"
    if "select" in joined and POSITIVE_IN_UNIT_RE.search(joined):
        return "Select units may have in-unit washer/dryer"
    if "laundry room" in joined or "facility" in joined or "facilities" in joined or "community" in joined:
        return "Shared laundry"
    return summarize_signal(contexts_, "Laundry mentioned; verify type")


def summarize_private_bath(contexts_: list[str], matched_target: str) -> str:
    if not contexts_:
        return "Unknown"
    joined = " ".join(contexts_).lower()
    if re.search(r"\b(private bath|private bathroom|private ensuite|en[- ]?suite|own bathroom|own bath)\b", joined):
        return "Yes / likely private bathroom"
    if "shared bath" in joined or "shared bathroom" in joined:
        return "No / shared bathroom mentioned"
    parsed = parse_unit_label(matched_target.split(",")[0]) if matched_target != "Unknown" else None
    if parsed and parsed["beds"] and parsed["baths"] and parsed["beds"] == parsed["baths"] and parsed["beds"] >= 2:
        return "Likely yes based on bed/bath count"
    return summarize_signal(contexts_, "Bathroom mentioned; verify private/shared")


def summarize_phone(text: str, call_contexts: list[str]) -> str:
    search_text = " ".join(call_contexts) if call_contexts else text[:5000]
    phones = unique_keep_order(PHONE_RE.findall(search_text))
    if not phones and CALL_RE.search(text):
        phones = unique_keep_order(PHONE_RE.findall(text[:12000]))
    return ", ".join(phones[:3]) if phones else "Unknown"


def summarize_signal(contexts_: list[str], default: str) -> str:
    if not contexts_:
        return default
    return contexts_[0][:220]


def compact_evidence(snippets: list[str]) -> str:
    return " || ".join(unique_keep_order(snippets)[:8])[:1800]


def notes_for_extraction(
    price_min: int | None,
    budget: int,
    matched_target: str,
    laundry_contexts: list[str],
    call_contexts: list[str],
) -> str:
    notes = []
    if price_min is None:
        notes.append("Price not found in public text.")
    elif price_min > budget:
        notes.append("Lowest extracted per-person price is over budget.")
    if matched_target == "Unknown":
        notes.append("Requested unit type not found.")
    if not laundry_contexts:
        notes.append("Laundry wording not found.")
    if call_contexts:
        notes.append("Call/contact wording found; use phone column if price needs confirmation.")
    return " ".join(notes)


def apply_overrides(result: ScreenResult, overrides: dict[str, object]) -> None:
    for key, value in overrides.items():
        if hasattr(result, key):
            setattr(result, key, value)
    if overrides and result.target_unit != "Unknown" and result.matched_target == "No":
        result.matched_target = "Yes"
    if overrides and result.price_min is not None and result.price_basis == "Unknown":
        result.price_basis = "manual / per person unless noted"
    if overrides:
        note = "Manual override applied."
        result.notes = f"{result.notes} {note}".strip()


def fit_status(result: ScreenResult, budget: int, private_bath_mode: str) -> str:
    has_unit = result.matched_target == "Yes"
    under_budget = result.price_min is not None and result.price_min <= budget
    has_in_unit = "in-unit" in result.laundry.lower() or "in home" in result.laundry.lower() or "in-home" in result.laundry.lower()
    private_ok = private_bath_mode != "require" or result.private_bath.lower().startswith("yes") or "likely yes" in result.private_bath.lower()
    if has_unit and under_budget and has_in_unit and private_ok:
        return "Strong fit"
    if has_unit and under_budget and private_ok:
        return "Budget fit; laundry tradeoff"
    if has_unit and under_budget and not private_ok:
        return "Budget fit; bathroom tradeoff"
    if has_unit and has_in_unit:
        return "Feature fit; price unknown/over"
    if has_unit:
        return "Target unit found; price/details need review"
    return "Needs review"


def confidence(result: ScreenResult) -> str:
    score = 0
    if result.matched_target == "Yes":
        score += 1
    if result.price_min is not None:
        score += 1
    if result.sqft_text != "Unknown":
        score += 1
    if result.laundry != "Unknown":
        score += 1
    if result.parking != "Unknown":
        score += 1
    if result.private_bath != "Unknown":
        score += 1
    if result.phone != "Unknown":
        score += 1
    if "Manual override applied" in result.notes:
        score += 2
    if score >= 5:
        return "High"
    if score >= 3:
        return "Medium"
    return "Low"


def sort_key(result: ScreenResult, budget: int, private_bath_mode: str) -> tuple[int, int, int, int, int, str]:
    has_unit = 0 if result.matched_target == "Yes" else 1
    under_budget = 0 if result.price_min is not None and result.price_min <= budget else 1
    private_bath = 0 if result.private_bath.lower().startswith("yes") or "likely yes" in result.private_bath.lower() else 1
    if private_bath_mode == "any":
        private_bath = 0
    in_unit = 0 if "in-unit" in result.laundry.lower() or "in-home" in result.laundry.lower() else 1
    price = result.price_min if result.price_min is not None else 999999
    return (has_unit, under_budget, private_bath, in_unit, price, result.name.lower())


def write_csv(path: str, results: list[ScreenResult]) -> None:
    fields = list(ScreenResult.__dataclass_fields__.keys())
    with open(path, "w", newline="", encoding="utf-8") as file:
        writer = csv.DictWriter(file, fieldnames=fields)
        writer.writeheader()
        for result in results:
            writer.writerow({field: getattr(result, field) for field in fields})


def write_xlsx(path: str, results: list[ScreenResult]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Apartment Screening"
    headers = [
        "Rank",
        "Apartment",
        "Group",
        "Address",
        "Distance meters",
        "Matching unit",
        "Matched target",
        "Price per person",
        "Price numeric per person",
        "Price basis",
        "Area",
        "Private bathroom",
        "Laundry",
        "Parking",
        "Pet policy",
        "Phone",
        "Fit status",
        "Confidence",
        "Notes",
        "Website",
        "Evidence",
        "Pages checked",
    ]
    ws.append(headers)
    for index, result in enumerate(results, start=1):
        ws.append(
            [
                index,
                result.name,
                result.group,
                result.address,
                result.distance_meters,
                result.target_unit,
                result.matched_target,
                result.price_text,
                result.price_min,
                result.price_basis,
                result.sqft_text,
                result.private_bath,
                result.laundry,
                result.parking,
                result.pet_policy,
                result.phone,
                result.fit_status,
                result.confidence,
                result.notes,
                result.website,
                result.evidence,
                result.pages_checked,
            ]
        )

    table = Table(displayName="ApartmentScreeningAuto", ref=f"A1:V{ws.max_row}")
    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True)
    ws.add_table(table)
    ws.freeze_panes = "A2"
    widths = {
        "A": 7,
        "B": 30,
        "C": 28,
        "D": 38,
        "E": 14,
        "F": 20,
        "G": 14,
        "H": 20,
        "I": 18,
        "J": 26,
        "K": 18,
        "L": 26,
        "M": 30,
        "N": 30,
        "O": 24,
        "P": 18,
        "Q": 30,
        "R": 12,
        "S": 38,
        "T": 42,
        "U": 60,
        "V": 60,
    }
    for column, width in widths.items():
        ws.column_dimensions[column].width = width
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
    for row in range(2, ws.max_row + 1):
        status = ws[f"Q{row}"].value
        fill = None
        if status == "Strong fit":
            fill = PatternFill("solid", fgColor="C6EFCE")
        elif status == "Budget fit; laundry tradeoff":
            fill = PatternFill("solid", fgColor="FFF2CC")
        elif status == "Budget fit; bathroom tradeoff":
            fill = PatternFill("solid", fgColor="FCE4D6")
        elif status == "Feature fit; price unknown/over":
            fill = PatternFill("solid", fgColor="D9EAF7")
        if fill:
            for cell in ws[row]:
                cell.fill = fill
    wb.save(path)


def write_docx(
    path: str,
    results: list[ScreenResult],
    budget: int,
    acceptable_units: list[dict[str, int | str | None]],
    private_bath_mode: str,
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run("Apartment Website Screening")
    run.bold = True
    run.font.size = Pt(18)

    target_labels = ", ".join(str(unit["label"]) for unit in acceptable_units)
    doc.add_paragraph(
        f"Target units: {target_labels}. Budget: <= ${budget:,} per person. "
        f"Private bathroom: {private_bath_mode}. In-unit laundry preferred. Pet policy is recorded but not ranked."
    )
    doc.add_heading("Top Candidates", level=1)
    shortlist = [result for result in results if result.fit_status in {"Strong fit", "Budget fit; laundry tradeoff", "Budget fit; bathroom tradeoff", "Feature fit; price unknown/over"}]
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    for index, header in enumerate(["Apartment", "Unit", "Price/person", "Basis", "Area", "Private bath", "Phone", "Status"]):
        table.rows[0].cells[index].text = header
    for result in shortlist[:18]:
        cells = table.add_row().cells
        values = [
            result.name,
            result.target_unit,
            result.price_text,
            result.price_basis,
            result.sqft_text,
            result.private_bath,
            result.phone,
            result.fit_status,
        ]
        for index, value in enumerate(values):
            cells[index].text = str(value)

    doc.add_heading("Notes", level=1)
    doc.add_paragraph("Automatically extracted website fields can miss content loaded only after clicking, embedded in images, or hidden behind availability forms.", style="List Bullet")
    doc.add_paragraph("Manual overrides are supported for confirmed leasing-office details or phone/email answers.", style="List Bullet")
    doc.add_paragraph("Rows marked Unknown or Needs review should be verified before applying.", style="List Bullet")
    doc.save(path)


def normalize_name(name: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", name.lower()).strip()


def normalize_url(url: str) -> str:
    parsed = urllib.parse.urlparse(url)
    return urllib.parse.urlunparse((parsed.scheme, parsed.netloc.lower(), parsed.path.rstrip("/"), "", parsed.query, ""))


def clean_snippet(value: str) -> str:
    value = re.sub(r"\s+", " ", value)
    return value.strip(" -|:;,")


def unique_keep_order(values: list[str]) -> list[str]:
    seen = set()
    result = []
    for value in values:
        key = value.strip()
        if not key or key in seen:
            continue
        seen.add(key)
        result.append(key)
    return result


if __name__ == "__main__":
    raise SystemExit(main())
