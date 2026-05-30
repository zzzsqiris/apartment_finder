#!/usr/bin/env python3
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Side, Border
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo


OUTPUT_XLSX = Path("ucla_apartment_screening.xlsx")
OUTPUT_DOCX = Path("ucla_apartment_screening_summary.docx")


RECORDS = [
    {
        "name": "Ariel Court",
        "group": "North Westwood / Strathmore-Gayley",
        "address": "535 Gayley Ave, Los Angeles, CA 90024",
        "unit": "Studio",
        "price": 2399,
        "price_text": "$2,399+",
        "sqft": "530+",
        "laundry": "Shared building laundry",
        "parking": "Covered parking",
        "pet_policy": "Pet-friendly / pets upon approval",
        "website": "https://www.decron.com/apartments/ca/los-angeles/ariel-court/floor-plans",
        "source": "Official floorplans + Zillow/RentCafe snippets",
        "confidence": "High",
        "notes": "Studio units under $2,500; 1BR starts higher. No in-unit laundry found.",
    },
    {
        "name": "Atrium",
        "group": "North Westwood / Strathmore-Gayley",
        "address": "10965 Strathmore Dr, Los Angeles, CA 90024",
        "unit": "1B1B",
        "price": 3500,
        "price_text": "$3,500+",
        "sqft": "650+",
        "laundry": "Unknown / not listed as in-unit",
        "parking": "Assigned garage parking / EV charging",
        "pet_policy": "No pets",
        "website": "https://www.atrium-westwood.com/floor-plans/",
        "source": "Official floorplans and FAQ",
        "confidence": "High",
        "notes": "Over budget for 1B1B.",
    },
    {
        "name": "Village Villa Apartments",
        "group": "North Westwood / Strathmore-Gayley",
        "address": "445 Landfair Ave, Los Angeles, CA 90024",
        "unit": "1B1B",
        "price": 2750,
        "price_text": "$2,750+",
        "sqft": "Unknown",
        "laundry": "Shared laundry room",
        "parking": "Underground parking, extra cost",
        "pet_policy": "Unknown",
        "website": "https://www.villagevillaapartments.com/",
        "source": "Official website post",
        "confidence": "Medium",
        "notes": "Website availability post appears dated; verify current pricing.",
    },
    {
        "name": "430 Kelton",
        "group": "Southwest Westwood",
        "address": "430 Kelton Ave, Los Angeles, CA 90024",
        "unit": "1B1B",
        "price": 3215,
        "price_text": "$3,215+ / call for lower plans",
        "sqft": "550-690",
        "laundry": "Shared laundry facility",
        "parking": "Assigned parking / garage",
        "pet_policy": "No pets",
        "website": "https://www.430kelton.com/floor-plans/",
        "source": "Official floorplans and amenities",
        "confidence": "High",
        "notes": "1B1B Plan C listed at $3,215; Plans A/B call for pricing.",
    },
    {
        "name": "Gayley Heights",
        "group": "North Westwood / Strathmore-Gayley",
        "address": "10995 Le Conte Ave, Los Angeles, CA 90024",
        "unit": "UCLA housing",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "https://portal.housing.ucla.edu/single-undergraduate-university-apartments",
        "source": "UCLA housing portal",
        "confidence": "Low",
        "notes": "University apartment; likely eligibility-specific, not a normal market-rate option.",
    },
    {
        "name": "Westwood Executive House",
        "group": "Southwest Westwood",
        "address": "424 Kelton Ave, Los Angeles, CA 90024",
        "unit": "1B1B",
        "price": 2775,
        "price_text": "$2,775+",
        "sqft": "640",
        "laundry": "Shared building laundry",
        "parking": "1 free parking spot / covered parking",
        "pet_policy": "No pets",
        "website": "https://www.apartments.com/westwood-executive-house-los-angeles-ca/yq6sn4s/",
        "source": "Apartments.com / Zillow snippets",
        "confidence": "High",
        "notes": "Over $2,500; parking is a plus.",
    },
    {
        "name": "Midvale Apartments",
        "group": "Southwest Westwood",
        "address": "527 Midvale Ave, Los Angeles, CA 90024",
        "unit": "Studio",
        "price": 2600,
        "price_text": "$2,600+",
        "sqft": "500-529",
        "laundry": "Shared laundry",
        "parking": "Parking single $150/mo; tandem options listed",
        "pet_policy": "Unknown",
        "website": "https://www.midvaleapts.com/floor-plans/",
        "source": "Official fee page + Zillow/Rent snippets",
        "confidence": "High",
        "notes": "Close to budget but above $2,500 in current listings.",
    },
    {
        "name": "Village Lofts",
        "group": "North Westwood / Strathmore-Gayley",
        "address": "11024 Strathmore Dr, Los Angeles, CA 90024",
        "unit": "Unknown",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Possibly in-unit in some units",
        "parking": "Optional parking noted in sublease posts",
        "pet_policy": "Unknown",
        "website": "http://villageloftsla.com/",
        "source": "Website + recent sublease snippets",
        "confidence": "Low",
        "notes": "Could be worth asking, but official pricing was not found.",
    },
    {
        "name": "11017 Strathmore",
        "group": "North Westwood / Strathmore-Gayley",
        "address": "11017 Strathmore Dr, Los Angeles, CA 90024",
        "unit": "Unknown",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "https://www.apartments.com/11017-strathmore-dr-los-angeles-ca/xgxrf7m/",
        "source": "Apartments.com listing",
        "confidence": "Low",
        "notes": "Not enough public data found.",
    },
    {
        "name": "Westfield Residence",
        "group": "Southwest Westwood",
        "address": "500 Kelton Ave, Los Angeles, CA 90024",
        "unit": "Room/board housing",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Room",
        "laundry": "Shared laundry facilities",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "http://www.westfieldresidence.org/housing",
        "source": "Official housing page",
        "confidence": "High",
        "notes": "Women-only room-and-board style housing; not a Studio/1B1B apartment.",
    },
    {
        "name": "520 Kelton",
        "group": "Southwest Westwood",
        "address": "520 Kelton Ave, Los Angeles, CA 90024",
        "unit": "1B1B",
        "price": None,
        "price_text": "Call for details",
        "sqft": "634",
        "laundry": "Unknown",
        "parking": "Parking varies by unit",
        "pet_policy": "Pet-friendly",
        "website": "https://www.keltonapartments.com/floorplans/1-bedroom%2C-1-bath",
        "source": "Official RentCafe floorplan + UniShack snippets",
        "confidence": "Medium",
        "notes": "Good to call because price is not publicly shown.",
    },
    {
        "name": "Kelton Arms",
        "group": "Southwest Westwood",
        "address": "433 Kelton Ave, Los Angeles, CA 90024",
        "unit": "1B1B",
        "price": 3500,
        "price_text": "$3,500",
        "sqft": "Unknown / large",
        "laundry": "Shared on-site laundry",
        "parking": "Assigned underground parking",
        "pet_policy": "No pets",
        "website": "https://www.keltonarms.com/floorplans/",
        "source": "Official site + CollegeRentals listing",
        "confidence": "Medium",
        "notes": "Over budget; some third-party text mentioned washer/dryer in unit but official amenity page says laundry facility.",
    },
    {
        "name": "Axiom Westwood",
        "group": "North Westwood / Strathmore-Gayley",
        "address": "415 Gayley Ave, Los Angeles, CA 90024",
        "unit": "Studio",
        "price": 2353,
        "price_text": "$2,353-$2,403",
        "sqft": "525",
        "laundry": "Shared laundry care center",
        "parking": "Covered/shared garage parking",
        "pet_policy": "Pets allowed",
        "website": "https://www.axiomwestwood.com/brochure",
        "source": "Official brochure + Rentable listing",
        "confidence": "High",
        "notes": "Under $2,500 for studio, but no in-unit laundry found.",
    },
    {
        "name": "Tiverton Court",
        "group": "Westwood Village Core",
        "address": "940 Tiverton Ave, Los Angeles, CA 90024",
        "unit": "Studio / 1B1B",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Laundry rooms every floor; select units have in-unit laundry",
        "parking": "Gated assigned parking",
        "pet_policy": "Unknown",
        "website": "https://www.tivertoncourt.com/",
        "source": "Official website and flyer",
        "confidence": "High",
        "notes": "Worth contacting because select apartments have in-unit laundry.",
    },
    {
        "name": "Strathmore Regency Apartments",
        "group": "North Westwood / Strathmore-Gayley",
        "address": "11050 Strathmore Dr, Los Angeles, CA 90024",
        "unit": "Unknown",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "https://www.strathmoreregency.com/",
        "source": "Official website",
        "confidence": "Low",
        "notes": "Public availability not found in this pass.",
    },
    {
        "name": "Tipuana Apartments",
        "group": "Westwood Village / Weyburn-Levering",
        "address": "900 Weyburn Pl, Los Angeles, CA 90024",
        "unit": "UCLA housing",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "https://portal.housing.ucla.edu/",
        "source": "UCLA housing portal",
        "confidence": "Low",
        "notes": "University apartment; likely eligibility-specific.",
    },
    {
        "name": "The Glendon Apartments in Westwood",
        "group": "Westwood Village Core",
        "address": "1040 Glendon Ave, Los Angeles, CA 90024",
        "unit": "1B1B",
        "price": None,
        "price_text": "Unknown / likely premium",
        "sqft": "Unknown",
        "laundry": "In-unit washer/dryer",
        "parking": "Parking garage",
        "pet_policy": "Dog friendly / pet-friendly noted",
        "website": "https://douglasemmettapartments.com/westwood-apartments/the-glendon/",
        "source": "VeryApt / housing guide snippets",
        "confidence": "Medium",
        "notes": "Strong laundry fit, but pricing not found and likely above budget.",
    },
    {
        "name": "Strathmore Veteran Apartments",
        "group": "North Westwood / Strathmore-Gayley",
        "address": "11090 Strathmore Dr, Los Angeles, CA 90024",
        "unit": "1B1B",
        "price": 2694,
        "price_text": "$2,694 current official; older sublease $2,495",
        "sqft": "450-550",
        "laundry": "Unknown",
        "parking": "Free parking / 1 parking spot noted",
        "pet_policy": "Unknown",
        "website": "https://www.liveatstrathmoreveteran.com/availableunits",
        "source": "Official available units + Reddit sublease snippets",
        "confidence": "Medium",
        "notes": "Slightly over budget in official listing; old sublease data was under $2,500.",
    },
    {
        "name": "Strathmore Arms",
        "group": "North Westwood / Strathmore-Gayley",
        "address": "11090 Strathmore Dr, Los Angeles, CA 90024",
        "unit": "1B1B",
        "price": None,
        "price_text": "Call / same portfolio",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Free parking noted for portfolio",
        "pet_policy": "Unknown",
        "website": "https://www.liveatstrathmoreveteran.com/",
        "source": "Official portfolio site",
        "confidence": "Low",
        "notes": "Likely same leasing portfolio as Strathmore Veteran.",
    },
    {
        "name": "Levering Apartments",
        "group": "Westwood Village / Weyburn-Levering",
        "address": "628 Levering Ave, Los Angeles, CA 90024",
        "unit": "Unknown",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "https://www.leveringapts.com/",
        "source": "Official site",
        "confidence": "Low",
        "notes": "Could not verify floorplan/pricing in this pass.",
    },
    {
        "name": "El Greco Lofts",
        "group": "Westwood Village Core",
        "address": "1030 Tiverton Ave, Los Angeles, CA 90024",
        "unit": "Studio / 1B1B",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Single space parking $150/mo noted for related listing",
        "pet_policy": "No pets",
        "website": "https://www.elgrecolofts.com/",
        "source": "Official site + search snippets",
        "confidence": "Medium",
        "notes": "Studio/1BR layouts, but public price unavailable.",
    },
    {
        "name": "Westwood Village Apartments",
        "group": "Westwood Village Core",
        "address": "1033 Hilgard Ave, Los Angeles, CA 90024",
        "unit": "1B1B",
        "price": 2995,
        "price_text": "$2,995+",
        "sqft": "650-750",
        "laundry": "Laundry facilities / dry cleaning",
        "parking": "Unknown",
        "pet_policy": "Pet-friendly",
        "website": "https://www.westwoodvillageliving.com/floorplans",
        "source": "Official floorplans + Rent/Apartments.com snippets",
        "confidence": "High",
        "notes": "Over budget.",
    },
    {
        "name": "Lindbrook Manor",
        "group": "Westwood Village Core",
        "address": "10824 Lindbrook Dr, Los Angeles, CA 90024",
        "unit": "Studio",
        "price": 2098,
        "price_text": "$2,098 studio; Junior-1 starts $2,457",
        "sqft": "217 studio; 463-488 junior",
        "laundry": "Unknown / likely shared laundry room",
        "parking": "Unknown",
        "pet_policy": "Pet-friendly",
        "website": "https://www.lindbrookmanor.com/floorplans",
        "source": "Official floorplans",
        "confidence": "High",
        "notes": "Strong budget fit, but in-unit laundry not confirmed.",
    },
    {
        "name": "Gayley + Lindbrook Apartments",
        "group": "North Westwood / Strathmore-Gayley",
        "address": "1122 Gayley Ave, Los Angeles, CA 90024",
        "unit": "Studio / 1B1B",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Studio 670-785; 1B1B 703-1,176",
        "laundry": "In-unit washer/dryer",
        "parking": "Underground parking structure",
        "pet_policy": "Unknown",
        "website": "https://www.gayleyandlindbrook.com/floor-plans",
        "source": "Official floorplans/gallery + Daily Bruin",
        "confidence": "High",
        "notes": "Excellent feature fit; price not public, likely premium.",
    },
    {
        "name": "Legacy at Westwood Apartments",
        "group": "Wilshire Corridor",
        "address": "10833 Wilshire Blvd, Los Angeles, CA 90024",
        "unit": "1B1B",
        "price": None,
        "price_text": "Call for details",
        "sqft": "693+",
        "laundry": "In-unit washer/dryer",
        "parking": "3-story subterranean parking / guest parking / EV",
        "pet_policy": "Pet-friendly",
        "website": "https://www.legacyatwestwood.com/brochure.aspx",
        "source": "Official brochure / Greystar snippets",
        "confidence": "Medium",
        "notes": "Strong features, but price not shown.",
    },
    {
        "name": "Venezia",
        "group": "Wilshire Corridor",
        "address": "10795 Wilshire Blvd, Los Angeles, CA 90024",
        "unit": "Condo",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "http://www.toplacondos.com/Venezia,-The",
        "source": "Condo listing site",
        "confidence": "Low",
        "notes": "Condo building; rental availability not verified.",
    },
    {
        "name": "Marie Antoinette Condominiums",
        "group": "Wilshire Corridor",
        "address": "10787 Wilshire Blvd #1101, Los Angeles, CA 90024",
        "unit": "Condo",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "http://www.toplacondos.com/Marie-Antoinette",
        "source": "Condo listing site",
        "confidence": "Low",
        "notes": "Condo building; rental availability not verified.",
    },
    {
        "name": "Wilshire Holmby Condominiums",
        "group": "Wilshire Corridor",
        "address": "10433 Wilshire Blvd, Los Angeles, CA 90024",
        "unit": "Condo",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "https://www.highrises.com/los-angeles/the-wilshire-holmby/",
        "source": "Highrises listing",
        "confidence": "Low",
        "notes": "Condo building; rental availability not verified.",
    },
    {
        "name": "Wilshire Victoria Westwood Apartments",
        "group": "Wilshire Corridor",
        "address": "10700 Wilshire Ave, Los Angeles, CA 90024",
        "unit": "Unknown",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "http://www.wilshirevictoria.com/",
        "source": "Official website",
        "confidence": "Low",
        "notes": "Public pricing not found.",
    },
    {
        "name": "Wilshire Margot",
        "group": "Wilshire Corridor",
        "address": "10599 Wilshire Blvd, Los Angeles, CA 90024",
        "unit": "Junior 1 / Studio",
        "price": 3945,
        "price_text": "$3,945",
        "sqft": "535-609",
        "laundry": "In-home washer/dryer in select suites",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "https://www.wilshiremargot.com/floorplans/unfurnished-jr.-one-bedroom-one-bathroom",
        "source": "Official floorplan",
        "confidence": "High",
        "notes": "Over budget.",
    },
    {
        "name": "Wilshire Westwood Luxury Apartments",
        "group": "Wilshire Corridor",
        "address": "10530 Wilshire Blvd, Los Angeles, CA 90024",
        "unit": "Unknown",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "http://www.therobertsco.com/",
        "source": "Official/management site",
        "confidence": "Low",
        "notes": "Public pricing not found.",
    },
    {
        "name": "Westwood Riviera Apartments",
        "group": "Southwest Westwood",
        "address": "10969 Wellworth Ave, Los Angeles, CA 90024",
        "unit": "Studio",
        "price": 2856,
        "price_text": "$2,856+",
        "sqft": "500",
        "laundry": "Common laundry area per floor",
        "parking": "Garage noted",
        "pet_policy": "Pet-friendly",
        "website": "https://www.westwoodriviera.com/floorplans",
        "source": "Official floorplans",
        "confidence": "High",
        "notes": "Above $2,500.",
    },
    {
        "name": "Utama Royale Apartments",
        "group": "Wilshire Corridor",
        "address": "10351 Wilshire Blvd, Los Angeles, CA 90024",
        "unit": "Unknown",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "https://www.utamaroyale.com/",
        "source": "Official site",
        "confidence": "Low",
        "notes": "Public pricing not found.",
    },
    {
        "name": "The Rochester Apartments",
        "group": "Southwest Westwood",
        "address": "10959 Rochester Ave, Los Angeles, CA 90024",
        "unit": "Unknown",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "https://www.rochester-westwood.com/",
        "source": "Official site",
        "confidence": "Low",
        "notes": "Public pricing not found.",
    },
    {
        "name": "Midvale Court Apartments Westwood",
        "group": "Southwest Westwood",
        "address": "1400 Midvale Ave, Los Angeles, CA 90024",
        "unit": "Studio",
        "price": 1875,
        "price_text": "$1,875+ bachelor/studio",
        "sqft": "324",
        "laundry": "Shared laundry facilities",
        "parking": "Parking garage noted",
        "pet_policy": "Unknown",
        "website": "https://www.midvalecourt.com/floorplans",
        "source": "Official floorplans + ApartmentHomeLiving/VeryApt snippets",
        "confidence": "High",
        "notes": "Best budget match found, but very small studio and no in-unit laundry.",
    },
    {
        "name": "Westwood Kelton Towers",
        "group": "Southwest Westwood",
        "address": "1395 Kelton Ave, Los Angeles, CA 90024",
        "unit": "Unknown",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "http://westwoodkeltontowers.bhprop.com/",
        "source": "Official/management site",
        "confidence": "Low",
        "notes": "Public pricing not found.",
    },
    {
        "name": "Villa Bel Air",
        "group": "Other Westwood",
        "address": "130 S Sepulveda Blvd, Los Angeles, CA 90049",
        "unit": "Unknown",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "https://www.villabelairliving.com/",
        "source": "Official site",
        "confidence": "Low",
        "notes": "Farther west; public pricing not found.",
    },
    {
        "name": "Veteran Apartments",
        "group": "Southwest Westwood",
        "address": "1417 Veteran Ave, Los Angeles, CA 90024",
        "unit": "Unknown",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "https://www.veteranavenue.com/",
        "source": "Official site",
        "confidence": "Low",
        "notes": "Public pricing not found.",
    },
    {
        "name": "Westwood Village Galleria",
        "group": "Southwest Westwood",
        "address": "1441 Veteran Ave, Los Angeles, CA 90024",
        "unit": "Unknown",
        "price": None,
        "price_text": "Unknown",
        "sqft": "Unknown",
        "laundry": "Unknown",
        "parking": "Unknown",
        "pet_policy": "Unknown",
        "website": "http://westwoodvillagegalleria.bhprop.com/",
        "source": "Official/management site",
        "confidence": "Low",
        "notes": "Public pricing not found.",
    },
]


def has_target_unit(record):
    unit = record["unit"].lower()
    return any(token in unit for token in ["studio", "1b1b", "1 bed", "junior"])


def under_budget(record):
    return record["price"] is not None and record["price"] <= 2500


def in_unit_laundry(record):
    laundry = record["laundry"].lower()
    negative_phrases = [
        "not listed as in-unit",
        "no in-unit",
        "not in-unit",
        "unknown / not listed",
    ]
    if any(phrase in laundry for phrase in negative_phrases):
        return False
    return "in-unit" in laundry or "in-home" in laundry or "in suite" in laundry


def sort_score(record):
    price = record["price"] if record["price"] is not None else 999999
    return (
        0 if has_target_unit(record) else 1,
        0 if under_budget(record) else 1,
        0 if in_unit_laundry(record) else 1,
        price,
        record["name"].lower(),
    )


def fit_status(record):
    if not has_target_unit(record):
        return "No target unit confirmed"
    if under_budget(record) and in_unit_laundry(record):
        return "Strong fit"
    if under_budget(record):
        return "Budget fit; laundry tradeoff"
    if in_unit_laundry(record):
        return "Feature fit; price unknown/over"
    return "Does not meet full criteria"


def build_workbook(records):
    wb = Workbook()
    ws = wb.active
    ws.title = "Apartment Screening"

    headers = [
        "Rank",
        "Apartment",
        "Group",
        "Address",
        "Best matching unit",
        "Price",
        "Price numeric",
        "Area",
        "Laundry",
        "Parking",
        "Pet policy",
        "Fit status",
        "Data confidence",
        "Notes",
        "Website",
        "Source basis",
    ]
    ws.append(headers)
    sorted_records = sorted(records, key=sort_score)
    for idx, record in enumerate(sorted_records, 1):
        ws.append(
            [
                idx,
                record["name"],
                record["group"],
                record["address"],
                record["unit"],
                record["price_text"],
                record["price"],
                record["sqft"],
                record["laundry"],
                record["parking"],
                record["pet_policy"],
                fit_status(record),
                record["confidence"],
                record["notes"],
                record["website"],
                record["source"],
            ]
        )

    table = Table(displayName="ApartmentScreening", ref=f"A1:P{ws.max_row}")
    style = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True, showColumnStripes=False)
    table.tableStyleInfo = style
    ws.add_table(table)

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:P{ws.max_row}"
    widths = {
        "A": 8,
        "B": 30,
        "C": 28,
        "D": 42,
        "E": 20,
        "F": 22,
        "G": 14,
        "H": 16,
        "I": 32,
        "J": 32,
        "K": 24,
        "L": 28,
        "M": 16,
        "N": 48,
        "O": 48,
        "P": 34,
    }
    for col, width in widths.items():
        ws.column_dimensions[col].width = width
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.column_dimensions["G"].hidden = True
    ws.row_dimensions[1].height = 36

    green = PatternFill("solid", fgColor="E2F0D9")
    yellow = PatternFill("solid", fgColor="FFF2CC")
    blue = PatternFill("solid", fgColor="DDEBF7")
    for row in range(2, ws.max_row + 1):
        status = ws[f"L{row}"].value
        fill = None
        if status == "Strong fit":
            fill = green
        elif status == "Budget fit; laundry tradeoff":
            fill = yellow
        elif status == "Feature fit; price unknown/over":
            fill = blue
        if fill:
            for col in range(1, ws.max_column + 1):
                ws.cell(row=row, column=col).fill = fill

    summary = wb.create_sheet("Summary")
    summary_rows = [
        ["Generated", datetime.now().strftime("%Y-%m-%d %H:%M")],
        ["Source list", "apartments_google_budget_clean.csv; only properties with website retained"],
        ["User target", "Studio or 1B1B, budget <= $2,500, ideally in-unit washer/dryer"],
        ["Pet requirement", "Ignored per latest instruction"],
        ["Records reviewed", len(records)],
        ["Strong fits", sum(1 for r in records if fit_status(r) == "Strong fit")],
        ["Budget fits with laundry tradeoff", sum(1 for r in records if fit_status(r) == "Budget fit; laundry tradeoff")],
        ["Feature fits but price unknown/over", sum(1 for r in records if fit_status(r) == "Feature fit; price unknown/over")],
    ]
    for row in summary_rows:
        summary.append(row)
    summary.column_dimensions["A"].width = 34
    summary.column_dimensions["B"].width = 90
    for row in summary.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for cell in summary["A"]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="D9EAF7")

    wb.save(OUTPUT_XLSX)


def add_table(document, rows, columns):
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, col in enumerate(columns):
        hdr[idx].text = col
        for p in hdr[idx].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            cells[idx].text = str(row.get(col, ""))
    return table


def build_doc(records):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for style_name in ["Heading 1", "Heading 2"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.color.rgb = None

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("UCLA Apartment Screening Summary")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} | Studio/1B1B, <= $2,500 target, no pet ranking").italic = True

    doc.add_heading("Executive Takeaways", level=1)
    takeaways = [
        "No property in the reviewed set clearly satisfies all three original preferences at once: Studio/1B1B, under $2,500, and confirmed in-unit laundry.",
        "Best budget matches are Midvale Court, Lindbrook Manor, Axiom Westwood, and Ariel Court; all appear to have shared or unconfirmed laundry rather than confirmed in-unit laundry.",
        "Best in-unit laundry leads are Gayley + Lindbrook, The Glendon, Legacy at Westwood, and Wilshire Margot, but public pricing is unknown or above budget.",
        "Parking varies widely. Several buildings offer assigned/garage parking, often at extra cost; Westwood Executive House currently advertises a parking promotion.",
    ]
    for item in takeaways:
        doc.add_paragraph(item, style="List Bullet")

    sorted_records = sorted(records, key=sort_score)
    focus = [r for r in sorted_records if fit_status(r) in {"Strong fit", "Budget fit; laundry tradeoff", "Feature fit; price unknown/over"}]
    doc.add_heading("Shortlist", level=1)
    rows = []
    for r in focus[:18]:
        rows.append(
            {
                "Apartment": r["name"],
                "Unit": r["unit"],
                "Price": r["price_text"],
                "Area": r["sqft"],
                "Laundry": r["laundry"],
                "Parking": r["parking"],
                "Status": fit_status(r),
            }
        )
    add_table(doc, rows, ["Apartment", "Unit", "Price", "Area", "Laundry", "Parking", "Status"])

    doc.add_heading("Best Under-Budget Leads", level=1)
    budget_rows = [
        r for r in sorted_records if under_budget(r) and has_target_unit(r)
    ]
    if budget_rows:
        for r in budget_rows:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(r["name"]).bold = True
            p.add_run(f": {r['unit']}, {r['price_text']}, {r['sqft']}; laundry: {r['laundry']}; parking: {r['parking']}.")
    else:
        doc.add_paragraph("No under-budget Studio/1B1B leads were confirmed.")

    doc.add_heading("Caveats", level=1)
    caveats = [
        "Prices and availability change quickly; treat all numbers as a screening snapshot and verify before applying.",
        "When a site only listed 'call for details' or hid availability behind a contact form, the table marks price as Unknown rather than guessing.",
        "Pet policy is included where found but not used for ranking, per latest instruction.",
        "Some UCLA-owned or eligibility-restricted housing remains in the spreadsheet but is marked with lower fit status.",
    ]
    for item in caveats:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Source Links", level=1)
    for r in sorted_records:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(r["name"]).bold = True
        p.add_run(f": {r['website']}")

    doc.save(OUTPUT_DOCX)


def main():
    build_workbook(RECORDS)
    build_doc(RECORDS)
    print(f"Wrote {OUTPUT_XLSX}")
    print(f"Wrote {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
